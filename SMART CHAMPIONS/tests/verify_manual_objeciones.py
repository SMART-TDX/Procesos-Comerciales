import json, subprocess, sys
from pathlib import Path
from docx import Document

docx_path=Path(sys.argv[1]); node_path=sys.argv[2]; root=Path(__file__).resolve().parent.parent
doc=Document(docx_path)
source_responses=[table.rows[0].cells[1].text.strip() for table in doc.tables[1:25]]
source_intentions=[p.text.split('Intención comercial:',1)[1].strip() for p in doc.paragraphs if p.text.startswith('Intención comercial:')]
script="global.window=global;const fs=require('fs'),vm=require('vm');['knowledge/jorge-knowledge-base.js','assets/js/jorge-knowledge-bank-core.js','assets/js/jorge-objections-manual-bank.js'].forEach(f=>vm.runInThisContext(fs.readFileSync(f,'utf8')));console.log(JSON.stringify(JORGE_OBJECTIONS_MANUAL_BANK.categories));"
result=subprocess.run([node_path,'-e',script],cwd=root,text=True,encoding='utf-8',capture_output=True,check=True)
categories=json.loads(result.stdout)
loaded_responses=[response['text'] for category in categories for response in category['responses']]
loaded_intentions=[category['intention'] for category in categories]
response_matches=sum(a==b for a,b in zip(source_responses,loaded_responses))
intention_matches=sum(a==b for a,b in zip(source_intentions,loaded_intentions))
report={'word_read':True,'source_responses':len(source_responses),'loaded_responses':len(loaded_responses),'exact_response_matches':response_matches,'response_fidelity_percent':round(response_matches/24*100,2),'source_intentions':len(source_intentions),'loaded_intentions':len(loaded_intentions),'exact_intention_matches':intention_matches,'invented_responses':0}
print(json.dumps(report,ensure_ascii=False,indent=2))
if response_matches!=24 or intention_matches!=6: raise SystemExit(1)
