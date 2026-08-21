from __future__ import annotations

import json
import re
import sys
from pathlib import Path

import pdfplumber
from docx import Document
from openpyxl import load_workbook


def clean(value):
    if value is None:
        return ""
    text = str(value).replace("\u00a0", " ")
    return re.sub(r"[ \t]+", " ", text).strip()


def extract_pdf(path: Path):
    pages = []
    with pdfplumber.open(path) as pdf:
        for number, page in enumerate(pdf.pages, 1):
            text = clean(page.extract_text() or "")
            tables = []
            for table_index, table in enumerate(page.extract_tables() or [], 1):
                rows = [[clean(cell) for cell in row] for row in table if row]
                if rows:
                    tables.append({"table": table_index, "rows": rows})
            pages.append({"page": number, "text": text, "tables": tables})
    return {"type": "pdf", "pages": pages}


def extract_docx(path: Path):
    doc = Document(path)
    paragraphs = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    tables = []
    for table_index, table in enumerate(doc.tables, 1):
        rows = []
        for row_index, row in enumerate(table.rows, 1):
            values = [clean(cell.text) for cell in row.cells]
            if any(values):
                rows.append({"row": row_index, "values": values})
        if rows:
            tables.append({"table": table_index, "rows": rows})
    return {"type": "docx", "paragraphs": paragraphs, "tables": tables}


def extract_xlsx(path: Path):
    wb = load_workbook(path, data_only=False, read_only=True)
    sheets = []
    for ws in wb.worksheets:
        rows = []
        for row_index, row in enumerate(ws.iter_rows(values_only=False), 1):
            values = [clean(cell.value) for cell in row]
            if any(values):
                last = max(i for i, value in enumerate(values) if value)
                rows.append({"row": row_index, "values": values[: last + 1]})
        sheets.append({"sheet": ws.title, "rows": rows})
    return {"type": "xlsx", "sheets": sheets}


def main():
    output = Path(sys.argv[1])
    sources = [Path(value) for value in sys.argv[2:]]
    result = []
    for path in sources:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            data = extract_pdf(path)
        elif suffix == ".docx":
            data = extract_docx(path)
        elif suffix == ".xlsx":
            data = extract_xlsx(path)
        else:
            continue
        result.append({"path": str(path), "name": path.name, "data": data})
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"files": len(result), "output": str(output)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
